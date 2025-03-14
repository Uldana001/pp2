# from docx import Document

# filename = r"/Users/uldanakonyratbaeva/Desktop/opinion essay.docx"

# try:
#     doc = Document(filename)

#     text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

#     lines = [line for line in text.split("\n") if line.strip()]
#     num_lines = len(lines)

#     print(num_lines)
# except Exception as e:
#     print("Error: ", e)

from docx import Document

# File path
filename = r"/Users/uldanakonyratbaeva/Desktop/opinion essay.docx"

try:
    # Open the Word document
    doc = Document(filename)

    # Extract text from all paragraphs and handle internal line breaks
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    # Split by real line breaks inside paragraphs
    lines = text.splitlines()

    # Remove empty lines (if any)
    lines = [line for line in lines if line.strip()]

    # Count the number of lines
    num_lines = len(lines)

    print(f"Number of lines: {num_lines}")
except Exception as e:
    print("Error:", e)
